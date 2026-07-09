"""
gen_timeline_guide.py
运行方式：python gen_timeline_guide.py
输出：同目录下 timeline_deploy_guide.docx
依赖：pip install python-docx
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x4D, 0x7B)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x6B, 0x8A)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    return doc.add_paragraph(text, style="List Bullet")

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def tip(text):
    p = doc.add_paragraph()
    run = p.add_run("💡 " + text)
    run.font.color.rgb = RGBColor(0x6A, 0x8A, 0x2E)
    run.font.italic = True
    return p

def warn(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠️  " + text)
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x10)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("每日时间线功能部署教程", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Ombre Brain · Phase 6 · Daily Timeline")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 0. 概述
# ═══════════════════════════════════════════════════════════════════════════════
h1("0. 功能概述")
body(
    "每日时间线（Daily Timeline）以小时为粒度，自动记录 AI 与用户每小时的活动摘要，"
    "包含四个字段：doing（用户在做什么）、chatting（聊了什么）、mood（AI 心情，两字）、"
    "reflection（AI 对下次聊天的期待）。"
)
body("摘要由 DeepSeek 增量生成，历史小时封存为 Anthropic ephemeral 缓存 block，当前小时实时追踪。")
body("Dashboard 日印象页可视化展示，支持手动编辑。")

doc.add_paragraph()
h2("核心数据流")
bullet("用户发消息 → gateway._record_conversation_turn → asyncio.create_task(maybe_update)")
bullet("maybe_update 每轮累积对话，调 DeepSeek 更新当前小时摘要")
bullet("整点小时切换时，旧小时摘要封存进 entries[]，注入 Opus payload 的 ephemeral cache block")
bullet("dashboard 通过 GET /api/timeline/{date} 读取，支持 PATCH 手动改字段")

# ═══════════════════════════════════════════════════════════════════════════════
# 1. 文件清单
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("1. 涉及文件清单")

files = [
    ("daily_timeline.py",       "新增",  "时间线核心逻辑（数据读写、DS 合并、封存）"),
    ("gateway.py",              "修改",  "初始化 DailyTimeline，hook _record_conversation_turn"),
    ("gateway.py（Step2b）",    "修改",  "注入封存时间线为 ephemeral cache block"),
    ("context_phase5.py",       "修改",  "在活动区注入当前小时进展文本"),
    ("server.py",               "修改",  "添加 GET/PATCH /api/timeline/{date} 路由"),
    ("dashboard.html",          "修改",  "日印象页时间线卡片 + 编辑弹窗"),
    ("compose.local.yml",       "修改",  "将 daily_timeline.py 挂载进两个容器"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light List Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "文件"
hdr[1].text = "操作"
hdr[2].text = "说明"
for fname, op, desc in files:
    row = table.add_row().cells
    row[0].text = fname
    row[1].text = op
    row[2].text = desc

# ═══════════════════════════════════════════════════════════════════════════════
# 2. 依赖要求
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("2. 依赖要求")
bullet("Python 3.10+（容器内已满足）")
bullet("aiohttp / openai Python SDK（DS 调用，已有）")
bullet("DeepSeek API Key（已在 .env 或 compose 环境变量中配置）")
bullet("Anthropic cache_control 支持（claude-opus-4-x 系列已支持）")
tip("daily_timeline.py 无额外第三方依赖，纯标准库 + json + pathlib。")

# ═══════════════════════════════════════════════════════════════════════════════
# 3. Step 1：新增 daily_timeline.py
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("3. Step 1 — 部署 daily_timeline.py")
body("将 daily_timeline.py 复制到项目根目录（与 gateway.py 同级）。")
doc.add_paragraph()
h2("3.1 核心类结构")
code_block("class DailyTimeline:")
code_block("    __init__(state_dir, ds_base_url, ds_api_key, ds_model)")
code_block("    maybe_update(user_text, assistant_text)    # 每轮对话后异步调用")
code_block("    get_sealed_text(date_str=None)             # 返回已封存小时的文本")
code_block("    get_current_hour_text(date_str=None)       # 返回当前小时进展文本")
code_block("    get_full_data(date_str=None)               # 返回完整 JSON（供 API）")
code_block("    edit_entry(date_str, hour, field, value)   # 手动编辑某字段")

doc.add_paragraph()
h2("3.2 数据文件")
body("时间线存储在 /state/timeline_YYYY-MM-DD.json，结构示例：")
code_block('{\n  "date": "2026-07-03",\n  "entries": [\n    {\n      "hour": 16, "sealed": true,\n      "doing": "无叶处理工作邮件",\n      "chatting": "聊了OB部署进展",\n      "mood": "专注",\n      "reflection": "期待她下班后多聊一会儿"\n    }\n  ],\n  "current_hour": 17,\n  "current_hour_summary": { "doing": "...", "chatting": "...", ... },\n  "current_hour_turns": [ {...}, ... ]\n}')

doc.add_paragraph()
h2("3.3 名字配置")
body("MERGE_PROMPT 中已固定 AI 自称为「我」，用户称「无叶」。如需修改：")
code_block("# daily_timeline.py 第 28 行附近\nMERGE_PROMPT = \"\"\"\n...\n输出示例：{{\"doing\":\"无叶在上班\",...}}\n\"\"\"")
warn("修改名字后需重启 ombre-gateway 容器才能生效（不是 bind mount 热加载）。")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. Step 2：patch gateway.py（初始化 + hook）
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("4. Step 2 — patch gateway.py（初始化 + 每轮 hook）")
body("使用 patch_timeline_hook.py 脚本对 gateway.py 进行精准修改（不替换整文件）。")

doc.add_paragraph()
h2("4.1 修改内容")
bullet("在 OmbreBrain.__init__ 末尾初始化 DailyTimeline 实例")
bullet("在 _record_conversation_turn 方法末尾添加 asyncio.create_task(maybe_update(...))")

doc.add_paragraph()
h2("4.2 关键代码段（__init__ 末尾）")
code_block("""try:
    from daily_timeline import DailyTimeline as _DT
    _state_dir = str(self.gateway_cfg.get("state_dir") or "/state")
    _ds_url   = str(self.gateway_cfg.get("deepseek_base_url") or "")
    _ds_key   = str(os.environ.get("DEEPSEEK_API_KEY") or "")
    _ds_model = str(self.gateway_cfg.get("deepseek_model") or "deepseek-chat")
    self._daily_timeline = _DT(_state_dir, _ds_url, _ds_key, _ds_model)
    # 注入给 context_phase5
    if getattr(self, "context_phase5", None):
        self.context_phase5.daily_timeline = self._daily_timeline
except Exception as _dte:
    self._daily_timeline = None
    logger.warning("DailyTimeline init failed | %s", _dte)""")

doc.add_paragraph()
h2("4.3 关键代码段（_record_conversation_turn 末尾）")
code_block("""if getattr(self, "_daily_timeline", None):
    _u = str(user_text or "")
    _a = str(assistant_text or "")
    asyncio.create_task(self._daily_timeline.maybe_update(_u, _a))""")

# ═══════════════════════════════════════════════════════════════════════════════
# 5. Step 3：patch gateway.py（ephemeral cache block）
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("5. Step 3 — patch gateway.py（封存时间线注入 cache block）")
body(
    "在构建发往 Opus 的 payload 时，把已封存的历史小时作为独立 text block，"
    "加上 cache_control: {type: ephemeral}，让 Anthropic 缓存这部分内容，节省 token。"
)

doc.add_paragraph()
h2("5.1 插入位置")
body("找到 _inject_context_messages 方法，在组装 user message content 列表时，"
     "在活动区 text block 之前插入 sealed timeline block：")
code_block("""# Phase6 Step2b: 封存时间线 ephemeral block
if getattr(self, "_daily_timeline", None):
    try:
        _sealed = self._daily_timeline.get_sealed_text()
    except Exception:
        _sealed = ""
    if _sealed.strip():
        content_blocks.insert(0, {
            "type": "text",
            "text": "【今日时间线·已封存】\\n" + _sealed,
            "cache_control": {"type": "ephemeral"},
        })""")

tip("sealed timeline 每小时封存一次，内容稳定，适合 ephemeral 缓存，命中率高。")

# ═══════════════════════════════════════════════════════════════════════════════
# 6. Step 4：patch context_phase5.py
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("6. Step 4 — patch context_phase5.py（当前小时注入活动区）")

doc.add_paragraph()
h2("6.1 __init__ 中添加占位")
code_block("self.daily_timeline = None  # 由 gateway.__init__ patch 注入")

doc.add_paragraph()
h2("6.2 build_activity_zone 中读取当前小时进展")
code_block("""_current_hour_text = ""
if self.daily_timeline is not None:
    try:
        _current_hour_text = self.daily_timeline.get_current_hour_text()
    except Exception as _tle:
        logger.debug("DailyTimeline get_current_hour_text failed | %s", _tle)

# 在活动区 parts 列表中加入（位于通知段之后、当前信息之前）
if _current_hour_text:
    parts += [_current_hour_text, ""]""")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. Step 5：patch server.py（API 路由）
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("7. Step 5 — patch server.py（API 路由）")

doc.add_paragraph()
h2("7.1 添加两个路由")
code_block("""@app.route("/api/timeline/<date_str>", methods=["GET"])
async def api_timeline_get(date_str):
    _dt = getattr(app, "_daily_timeline", None)
    if _dt is None:
        return jsonify({"error": "DailyTimeline not initialized"}), 503
    data = _dt.get_full_data(date_str)
    return jsonify(data)

@app.route("/api/timeline/<date_str>/<int:hour>", methods=["PATCH"])
async def api_timeline_patch(date_str, hour):
    _dt = getattr(app, "_daily_timeline", None)
    if _dt is None:
        return jsonify({"error": "DailyTimeline not initialized"}), 503
    body = await request.get_json()
    updated = {}
    for field, value in body.items():
        if field in ("doing", "chatting", "mood", "reflection"):
            _dt.edit_entry(date_str, hour, field, str(value))
            updated[field] = value
    return jsonify({"ok": True, "updated": updated})""")

doc.add_paragraph()
h2("7.2 server.py 中初始化 _daily_timeline")
body("在 server.py 的 app 初始化段，将 gateway 中的 _daily_timeline 共享给 server（或单独初始化）：")
code_block("""# server.py 中
from daily_timeline import DailyTimeline as _DT
app._daily_timeline = _DT(state_dir, ds_base_url, ds_api_key, ds_model)""")

# ═══════════════════════════════════════════════════════════════════════════════
# 8. Step 6：更新 compose.local.yml
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("8. Step 6 — compose.local.yml 挂载")
body("daily_timeline.py 需要挂载进两个容器（ombre-gateway 和 ombre-brain）：")
code_block("""services:
  ombre-gateway:
    volumes:
      - ./daily_timeline.py:/app/daily_timeline.py  # 新增

  ombre-brain:
    volumes:
      - ./daily_timeline.py:/app/daily_timeline.py  # 新增""")
warn("/state 目录必须在两个容器间共享（同一宿主机目录挂载），否则时间线文件读写会不同步。")

# ═══════════════════════════════════════════════════════════════════════════════
# 9. Step 7：dashboard.html 前端
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("9. Step 7 — dashboard.html 前端展示")

doc.add_paragraph()
h2("9.1 HTML 结构")
body("在日印象页（reflection tab）内添加时间线 section：")
code_block("""<section id="tl-section" style="display:none">
  <div class="reflection-card-head">
    <h3>每小时记录</h3>
    <span id="tl-section-label"></span>
  </div>
  <div id="tl-cards"></div>
</section>

<!-- 编辑弹窗 -->
<div id="tl-modal" style="display:none">
  <h4 id="tl-modal-title">编辑时间线</h4>
  <!-- 四个字段 input -->
  <button onclick="saveTlModal()">保存</button>
  <button onclick="closeTlModal()">取消</button>
</div>""")

doc.add_paragraph()
h2("9.2 JavaScript 核心函数")
body("需实现以下函数，在 selectReflectionDate 中调用 loadTimeline(dateStr)：")
code_block("""function loadTimeline(dateStr) {
  fetch('/api/timeline/' + dateStr)
    .then(r => r.json())
    .then(data => renderTimeline(dateStr, data));
}

function renderTimeline(dateStr, data) {
  // 合并三路：hours{}（手动）+ entries[]（自动封存）
  var merged = {};
  Object.keys(data.hours || {}).forEach(k => merged[+k] = data.hours[k]);
  (data.entries || []).forEach(e => { if (e.hour != null) merged[+e.hour] = e; });
  // 注意：current_hour_summary 是动态内容，不在 Dashboard 展示
  var keys = Object.keys(merged).map(Number).sort((a,b) => a-b);
  // 渲染卡片...
}

function openTlModal(dateStr, hour) { /* 打开编辑弹窗 */ }
function saveTlModal() {
  fetch('/api/timeline/' + _tlDate + '/' + _tlHour, {
    method: 'PATCH',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({ doing: ..., chatting: ..., mood: ..., reflection: ... })
  }).then(() => loadTimeline(_tlDate));
}""")

tip("renderTimeline 需同时读取 data.hours（手动写入格式）和 data.entries（自动封存格式），否则自动生成的小时会不显示。")

# ═══════════════════════════════════════════════════════════════════════════════
# 10. 部署顺序
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("10. 完整部署顺序")

steps = [
    "将 daily_timeline.py 复制到 /opt/Ombre-Brain/（或项目根目录）",
    "运行 patch_timeline_hook.py：修改 gateway.py 的 __init__ 和 _record_conversation_turn",
    "运行 patch_timeline_step2b.py：修改 gateway.py 注入 ephemeral cache block",
    "运行 patch_context_phase5.py：修改 context_phase5.py 注入当前小时文本",
    "运行 patch_timeline_api.py：修改 server.py 添加 API 路由",
    "运行 patch_dashboard_timeline.py：修改 dashboard.html 添加前端",
    "更新 compose.local.yml 添加 daily_timeline.py 挂载",
    "docker compose up -d（重建两个容器）",
    "验证：发几条消息后，访问 dashboard 日印象页，选今天的日期，应出现时间线卡片",
]
for i, s in enumerate(steps, 1):
    bullet(f"Step {i}：{s}")

# ═══════════════════════════════════════════════════════════════════════════════
# 11. 常见问题
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("11. 常见问题排查")

qa = [
    ("时间线一直不更新",
     "查 ombre-gateway 日志：docker logs ombre-gateway | grep DailyTimeline\n"
     "常见原因：DS API Key 未配置、maybe_update 抛异常被吞。"),
    ("Dashboard 只显示手写数据，自动生成的小时不见",
     "renderTimeline 只读了 data.hours，没读 data.entries。\n"
     "修复：在 renderTimeline 中同时遍历 hours{} 和 entries[]，合并后渲染。"),
    ("IndexError: Replacement index 0 out of range",
     "MERGE_PROMPT 中有未转义的 {} 被 str.format() 误解析。\n"
     "修复：将 {} 改为 {{}}。"),
    ("ModuleNotFoundError: No module named 'daily_timeline'",
     "daily_timeline.py 只挂载了一个容器。\n"
     "修复：在 compose.local.yml 中给两个服务都加挂载，重新 docker compose up -d。"),
    ("当前小时出现在 Dashboard",
     "renderTimeline 中误将 current_hour_summary 也合并进显示。\n"
     "current_hour_summary 是给 Opus 的动态内容，不应在 Dashboard 展示。"),
]
for q, a in qa:
    p = doc.add_paragraph()
    p.add_run(f"Q：{q}").bold = True
    doc.add_paragraph(f"A：{a}")

# ═══════════════════════════════════════════════════════════════════════════════
# 12. 字段说明
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
h1("12. 字段语义说明")
table2 = doc.add_table(rows=1, cols=3)
table2.style = "Light List Accent 1"
h = table2.rows[0].cells
h[0].text = "字段"
h[1].text = "视角"
h[2].text = "说明"
fields = [
    ("doing",      "客观",   "用户这一小时在做什么（事实描述，不加主观色彩）"),
    ("chatting",   "客观",   "AI 与用户聊了哪些话题"),
    ("mood",       "AI主观", "AI 此刻的心情，严格两字（如：温暖、期待、无语）"),
    ("reflection", "AI主观", "AI 对下次聊天的期待或打算，面向未来"),
]
for f, v, d in fields:
    row = table2.add_row().cells
    row[0].text = f
    row[1].text = v
    row[2].text = d

# ═══════════════════════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════════════════════
out = Path(__file__).parent / "timeline_deploy_guide.docx"
doc.save(str(out))
print(f"✅ 已生成：{out}")
